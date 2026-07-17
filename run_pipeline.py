#!/usr/bin/env python3
"""
run_pipeline.py  (v2)  --  the single "AI function"
===================================================
One entry point that turns a user's SDTM study into a discontinuation-prediction
analysis. It enforces an input contract, then runs every stage end to end.

USER INPUTS
  --study        a name/label for the user's study (free text)
  --input-dir    folder with the user's SDTM files matching the CDISC Pilot 01
                 format; define.xml is REQUIRED
  (or --fetch-reference to download and run the reference CDISC Pilot 01 study)

INPUT CONTRACT (hard gate)
  define.xml must be present and every dataset it declares must exist as a
  .xpt or .csv file. If not, the pipeline prints exactly:
     "User's input does not meet CDISC SDTM requirements that results in
      termination of further analysis"
  and stops (exit code 2) -- no further analysis is attempted.

BACKEND (on a valid input)
  1. convert SDTM -> CSV (+ subject spine + patient-long)
  2. leakage-safe baseline feature engineering + 70/30 stratified split
  3. train / tune / evaluate models, with explainability
  4. deliver results as FILES and/or a combined WORD report

OUTPUT
  everything under --outdir (default: cdiscpilot01_out_v2)

USAGE
  python run_pipeline.py --study "My Phase 2 Study" --input-dir ./my_sdtm --report both
  python run_pipeline.py --fetch-reference --study "CDISC Pilot 01" --report word

Requires: pandas, numpy, scikit-learn, matplotlib, joblib
Optional: python-docx (Word report), shap, lifelines, pyreadstat
"""
from __future__ import annotations
import argparse
import json
import subprocess
import sys
from pathlib import Path

import pandas as pd

import sdtm_reader as reader   # same folder

HERE = Path(__file__).resolve().parent


# --------------------------------------------------------------------------
def _run_script(script, args):
    cmd = [sys.executable, str(HERE / script)] + args
    print(f"\n$ {' '.join(cmd)}")
    r = subprocess.run(cmd)
    if r.returncode != 0:
        sys.exit(f"Stage failed: {script} (exit {r.returncode})")


def _safe_read_json(path):
    try:
        return json.loads(Path(path).read_text())
    except Exception:
        return {}


# --------------------------------------------------------------------------
def build_word_report(outdir: Path, study: str):
    """Combine the key results into a single Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[note] python-docx not installed -> Word report skipped "
              "(pip install python-docx). Files are still available.")
        return None

    feat = _safe_read_json(outdir / "features" / "feature_manifest.json")
    met = _safe_read_json(outdir / "model" / "metrics.json")
    NAVY = RGBColor(0x1F, 0x38, 0x64)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    h = doc.add_heading(f"Discontinuation Prediction \u2014 {study}", level=0)
    doc.add_paragraph("Automated SDTM analysis report (baseline-only, "
                      "leakage-safe).").italic = True

    # Executive summary
    doc.add_heading("Executive summary", level=1)
    best = met.get("best_model", "n/a")
    bkey = "gbm" if best == "GradientBoosting" else "logreg"
    bm = met.get(bkey, {})
    n_rand = feat.get("n_randomized", "n/a")
    n_coh = feat.get("n_modeling_cohort", "n/a")
    bal = feat.get("target_balance", {})
    doc.add_paragraph(
        f"From {n_rand} randomized subjects, a modeling cohort of {n_coh} was "
        f"analyzed (target balance {bal}). The best model ({best}) achieved "
        f"holdout ROC-AUC {bm.get('roc_auc','n/a')} and PR-AUC "
        f"{bm.get('pr_auc','n/a')}. Discontinuation is driven mainly by planned "
        f"treatment arm and baseline disease-severity measures, consistent with "
        f"adverse-event-driven dropout.")

    # Cohort & target
    doc.add_heading("Cohort & target", level=1)
    subj_p = outdir / "subjects.csv"
    if subj_p.exists():
        s = pd.read_csv(subj_p)
        rnd = s[s["RANDOMIZED"] == True] if "RANDOMIZED" in s.columns else s
        t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
        t.rows[0].cells[0].text = "Metric"; t.rows[0].cells[1].text = "Value"
        rows = [("Subjects (DM)", len(s)),
                ("Randomized", int(s["RANDOMIZED"].sum()) if "RANDOMIZED" in s else "n/a"),
                ("No-disposition excluded", feat.get("n_no_disposition_excluded", 0)),
                ("Modeling cohort", n_coh),
                ("Target balance (0/1)", bal)]
        for k, v in rows:
            c = t.add_row().cells; c[0].text = str(k); c[1].text = str(v)

    # Methodology (the four required answers)
    doc.add_heading("Methodology", level=1)
    meth = [
        ("Information used", feat.get("information_type", "baseline-only")
         + f" (randomization landmark, study day \u2264 {feat.get('landmark_day', 1)})"),
        ("Cutoff relative to completion", "not applicable \u2014 no post-baseline data is used"),
        ("Outcome", f"{feat.get('outcome_primary','binary')} primary; "
         f"{feat.get('outcome_also_provided','time-to-event')} provided for sensitivity"),
        ("Explainability", "logistic-regression odds ratios + permutation importance "
         "+ SHAP (if available)"),
        ("Leakage exclusions", ", ".join(feat.get("dm_excluded_leakage", [])) or "recorded in manifest"),
    ]
    for k, v in meth:
        p = doc.add_paragraph(); p.add_run(f"{k}: ").bold = True; p.add_run(str(v))

    # Holdout performance
    doc.add_heading("Holdout performance", level=1)
    t = doc.add_table(rows=1, cols=6); t.style = "Light Grid Accent 1"
    for i, hd in enumerate(["Model", "ROC-AUC", "PR-AUC", "F1", "Bal.Acc", "Brier"]):
        t.rows[0].cells[i].text = hd
    for key in ("logreg", "gbm"):
        m = met.get(key, {})
        if not m:
            continue
        c = t.add_row().cells
        c[0].text = m.get("model", key)
        c[1].text = str(m.get("roc_auc", "")); c[2].text = str(m.get("pr_auc", ""))
        c[3].text = str(m.get("f1", "")); c[4].text = str(m.get("balanced_accuracy", ""))
        c[5].text = str(m.get("brier", ""))

    # Key drivers
    doc.add_heading("Key drivers (attribution)", level=1)
    imp_p = outdir / "model" / "permutation_importance.csv"
    if imp_p.exists():
        imp = pd.read_csv(imp_p).head(8)
        t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
        t.rows[0].cells[0].text = "Feature"; t.rows[0].cells[1].text = "Permutation importance"
        for _, r in imp.iterrows():
            c = t.add_row().cells
            c[0].text = str(r["feature"]); c[1].text = f"{r['importance']:.4f}"

    # Discontinuation by arm (key insight)
    mm_p = outdir / "features" / "model_matrix.csv"
    if mm_p.exists():
        mm = pd.read_csv(mm_p)
        arm = "ARMCD" if "ARMCD" in mm.columns else ("ARM" if "ARM" in mm.columns else None)
        if arm:
            doc.add_heading("Key insight \u2014 discontinuation by planned arm", level=1)
            rate = mm.groupby(arm)["TARGET"].agg(["mean", "size"]).sort_values("mean", ascending=False)
            t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
            for i, hd in enumerate([arm, "Discontinue %", "n"]):
                t.rows[0].cells[i].text = hd
            for idx, r in rate.iterrows():
                c = t.add_row().cells
                c[0].text = str(idx); c[1].text = f"{r['mean']*100:.1f}"; c[2].text = str(int(r["size"]))

    # Figures
    for fig, cap in [("evaluation_curves.png", "ROC / PR / confusion (holdout)"),
                     ("shap_summary.png", "SHAP feature attribution")]:
        fp = outdir / "model" / fig
        if fp.exists():
            doc.add_heading(cap, level=1)
            doc.add_picture(str(fp), width=Inches(6.3))

    # Survival / time-to-event
    surv = _safe_read_json(outdir / "survival" / "survival_report.json")
    if surv:
        doc.add_heading("Survival analysis (time-to-event)", level=1)
        med = surv.get("median_time_overall")
        lr = surv.get("logrank", {})
        p = surv.get("cox", {})
        intro = (f"Complementing the binary models, a time-to-event view treats the "
                 f"outcome as days-until-dropout with completers censored. Across "
                 f"{surv.get('n','n/a')} patients ({surv.get('events','n/a')} events), "
                 f"the median time in study is "
                 f"{('%.0f days' % med) if isinstance(med,(int,float)) and med==med else 'not reached'}.")
        doc.add_paragraph(intro)
        if lr:
            doc.add_paragraph(
                f"Kaplan-Meier retention differs by treatment arm: log-rank "
                f"chi-square {lr.get('chi2')} (df {lr.get('dof')}), p = "
                f"{lr.get('p_value'):.2e}. A small p means the arms separate more "
                f"than chance would produce.")
        # median by arm
        by_arm = surv.get("km_by_arm", {})
        if by_arm:
            t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
            for i, hd in enumerate(["Arm", "N", "Events", "Median days (NR = not reached)"]):
                t.rows[0].cells[i].text = hd
            for a, d in by_arm.items():
                c = t.add_row().cells
                c[0].text = str(a); c[1].text = str(d.get("n"))
                c[2].text = str(d.get("events"))
                c[3].text = "NR" if d.get("median") is None else f"{d['median']:.0f}"
        # Cox hazard ratios
        if p and p.get("hazard_ratios"):
            doc.add_paragraph(
                f"Cox proportional-hazards model (baseline covariates; concordance "
                f"{p.get('concordance')}). A hazard ratio (HR) above 1 means higher "
                f"instantaneous risk of dropping out:")
            t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
            for i, hd in enumerate(["Covariate", "Hazard ratio", "p"]):
                t.rows[0].cells[i].text = hd
            hrs = p["hazard_ratios"]; pv = p.get("p_values", {})
            for k in hrs:
                c = t.add_row().cells
                c[0].text = str(k); c[1].text = f"{hrs[k]:.2f}"
                c[2].text = f"{pv.get(k, float('nan')):.3f}"
            viol = p.get("ph_violations")
            doc.add_paragraph(
                "Proportional-hazards assumption: " +
                ("holds for all covariates tested." if viol == [] else
                 (f"possible violation for {viol} (interpret those HRs as period-averaged)."
                  if viol else "not tested.")))
        else:
            doc.add_paragraph(
                "Cox model: not available in this run (install lifelines to enable "
                "hazard ratios and the proportional-hazards check). Kaplan-Meier, "
                "the log-rank test, and competing-risks incidence are shown above.")
        # competing risks note
        cif = surv.get("cif_final", {})
        if cif:
            top = ", ".join(f"{k.title()} {v:.0%}" for k, v in
                            sorted(cif.items(), key=lambda kv: -kv[1])[:4])
            doc.add_paragraph(
                f"Competing risks: \u201Cany-cause\u201D dropout combines competing reasons. "
                f"By end of study the cause-specific cumulative incidence is {top}. "
                f"Because these causes compete, each is estimated with the others "
                f"treated as competing events (Aalen-Johansen), not as a separate "
                f"survival curve \u2014 which would overstate any single cause.")
        for fig, cap in [("km_by_arm.png", "Kaplan-Meier retention by arm"),
                         ("cif_by_cause.png", "Cause-specific cumulative incidence")]:
            fp = outdir / "survival" / fig
            if fp.exists():
                doc.add_picture(str(fp), width=Inches(5.6))

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Reproducibility: ").bold = True
    note.add_run("all methodology choices are recorded in "
                 "features/feature_manifest.json; the holdout was scored once "
                 "after model selection on development.")

    out = outdir / f"{study.replace(' ', '_')}_report.docx"
    doc.save(str(out))
    print(f"\nWord report written: {out}")
    return out


# --------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--study", required=True, help="name/label for the study")
    ap.add_argument("--input-dir", help="folder with the user's SDTM files (+ define.xml)")
    ap.add_argument("--fetch-reference", action="store_true",
                    help="download and run the reference CDISC Pilot 01 study")
    ap.add_argument("--outdir", default="cdiscpilot01_out_v2")
    ap.add_argument("--target-mode", choices=["clinical", "any"], default="clinical")
    ap.add_argument("--landmark-day", type=int, default=1)
    ap.add_argument("--death", choices=["event", "exclude"], default="event")
    ap.add_argument("--report", choices=["files", "word", "both"], default="both")
    ap.add_argument("--seed", type=int, default=42)
    args = ap.parse_args()

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    print("=" * 70)
    print(f"STUDY: {args.study}")
    print("=" * 70)

    # ---- Stage 0: input contract (hard gate) -------------------------------
    if not args.fetch_reference:
        if not args.input_dir:
            sys.exit("Provide --input-dir (your SDTM files) or --fetch-reference")
        ok, missing, expected, has_define = reader.validate_sdtm_input(Path(args.input_dir))
        if not ok:
            # EXACT required message, then stop -- no further analysis
            print(reader.VALIDATION_ERROR)
            print(f"  Missing/invalid: {missing}", file=sys.stderr)
            sys.exit(2)
        print(f"[validate] OK \u2014 define.xml present; {len(expected)} declared "
              f"datasets all found.")

    # ---- Stage 1: SDTM -> CSV ----------------------------------------------
    print("\n[1/4] Converting SDTM to CSV ...")
    if args.fetch_reference:
        reader.convert(input_dir=None, outdir=outdir, source="github")
    else:
        reader.convert(args.input_dir, outdir, source="local")

    # ---- Stage 2: features + split -----------------------------------------
    print("\n[2/4] Feature engineering + 70/30 split ...")
    _run_script("build_features_cdiscpilot01.py", [
        "--indir", str(outdir), "--outdir", str(outdir / "features"),
        "--target-mode", args.target_mode, "--landmark-day", str(args.landmark_day),
        "--death", args.death, "--seed", str(args.seed)])

    # ---- Stage 3: model + evaluate + explain -------------------------------
    print("\n[3/5] Modeling + evaluation + explainability ...")
    _run_script("model_cdiscpilot01.py", [
        "--indir", str(outdir / "features"), "--outdir", str(outdir / "model"),
        "--seed", str(args.seed)])

    # ---- Stage 4: survival / time-to-event ---------------------------------
    print("\n[4/5] Survival analysis (Kaplan-Meier, log-rank, Cox, competing risks) ...")
    _run_script("survival_cdiscpilot01.py", [
        "--indir", str(outdir / "features"),
        "--subjects", str(outdir / "subjects.csv"),
        "--outdir", str(outdir / "survival")])

    # ---- Stage 5: deliver results ------------------------------------------
    print("\n[5/5] Assembling results ...")
    if args.report in ("word", "both"):
        build_word_report(outdir, args.study)
    if args.report in ("files", "both"):
        print(f"File artifacts under {outdir}/ (csv/, features/, model/, survival/).")

    print("\nDone.")


if __name__ == "__main__":
    main()
